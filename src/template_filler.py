"""
Template filler – opens a DOCX template and populates it with
extracted project data, classified device lists, and AI-generated content.

All six workpackage templates (WAN, Perimeter Firewall, OT Automation Machine,
MDF, IT Hardware, OT Hardware) share a standardised Deliverables section that
is aligned with the AOT milestone plan (MS1.3 → MS5).  The standard structure
is defined in STANDARD_DELIVERABLES below and is already embedded in each DOCX
template; it does not need to be re-injected at fill-time for normal use.
"""
from __future__ import annotations

import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt

# Default section number used when the deliverables heading carries no numeric prefix.
# Most templates number their Deliverables section as "4."; MDF/OT Automation use "5.".
# The actual value is always extracted from the heading text at runtime.
_DEFAULT_DELIVERABLES_SECTION = "4"

# ---------------------------------------------------------------------------
# Standard Deliverables – AOT Milestone-Aligned
# ---------------------------------------------------------------------------
# Each entry: (sub_heading_template, [bullet_text, …])
# {sec} is replaced with the document's deliverables section number (4 or 5).
STANDARD_DELIVERABLES: list[tuple[str, list[str]]] = [
    (
        "{sec}.1 MS1.3 – Provision of Design & Procurement",
        [
            "Procurement, design and delivery of all equipment as defined in Sections 1–3",
            "Physical installation of equipment in designated locations",
            "Installation of accessories and peripherals as required",
        ],
    ),
    (
        "{sec}.2 MS2 – FAT (Factory Acceptance Testing)",
        [
            "Factory acceptance testing and verification of all hardware prior to site delivery",
            "Pre-staging and configuration of equipment as per technical requirements",
            "Confirmation of full Bill of Materials receipt and manufacturer warranty coverage",
        ],
    ),
    (
        "{sec}.3 MS3 – SAT (Site Acceptance Testing & Commissioning)",
        [
            "Physical installation and commissioning of all equipment on site",
            "Connectivity and integration testing (network, power, and application layer)",
            "Validation that all delivered equipment meets Section 3 requirements",
            "All equipment must be validated and ready for operational use prior to Go-Live",
        ],
    ),
    (
        "{sec}.4 MS4 – UAT (User Acceptance Testing)",
        [
            "Operations user acceptance testing",
            "Asset inventory including: Device type, Model, Serial number",
            "Mapping of devices to: Users, Stations, Rooms or operational areas",
            "Handover of equipment to site management and operations team",
            "Provision of support and escalation contacts",
        ],
    ),
    (
        "{sec}.5 MS5 – Hypercare & Project Closure",
        [
            "Hypercare support during initial operational period",
            "Confirm NOC monitoring and SLA reports (where applicable)",
            "Validate all documentation required for handover to support is complete",
            "Final project acceptance sign-off and closure",
        ],
    ),
]


class TemplateFiller:
    """Fill DOCX workpackage templates with project-specific data."""

    def fill_template(
        self,
        template_path: str,
        template_type: str,
        project_info: dict[str, str],
        classified_devices: dict[str, list[dict]],
        generated_content: dict[str, Any],
    ) -> Document:
        """Open the template, fill it, and return the modified Document."""
        doc = Document(template_path)

        dispatch = {
            "IT Device": self._fill_it_device,
            "OT Device": self._fill_ot_device,
            "MDF": self._fill_mdf,
            "OT Automation Machine": self._fill_ot_automation,
        }
        filler_fn = dispatch.get(template_type, self._fill_generic)
        filler_fn(doc, project_info, classified_devices, generated_content)
        return doc

    # ── IT Device ─────────────────────────────────────────────────────────────

    def _fill_it_device(
        self,
        doc: Document,
        project_info: dict,
        classified_devices: dict,
        generated: dict,
    ) -> None:
        # Header fields
        self._replace_all(doc, "Project Name: ", f"Project Name: {project_info.get('project_name', '')}")
        self._replace_all(doc, "FBM ID:", f"FBM ID: {project_info.get('fbm_id', '')}")
        self._replace_all(doc, "Site address: ", f"Site address: {project_info.get('site_address', '')}")
        self._replace_all(doc, "SITE ID: ", f"SITE ID: {project_info.get('site_id', '')}")
        self._replace_all(
            doc,
            "xx/xx/2026 – Access to Site",
            f"{project_info.get('access_date', 'TBD')} – Access to Site",
        )
        self._replace_all(
            doc,
            "xx/xx/2026 – Go live.",
            f"{project_info.get('go_live_date', 'TBD')} – Go live.",
        )

        # Document Control table (Table 0, row 1, col 1) – Issue Date
        self._set_table_cell(doc, 0, 1, 1, datetime.today().strftime("%d-%m-%Y"))

        # Quantity Overview table (Table 1)
        it_devices = classified_devices.get("IT", [])
        admin_qty = str(generated.get("admin_users") or self._sum_matching(it_devices, ["laptop", "notebook"]))
        inbound_qty = str(generated.get("inbound_stations") or 0)
        pack_qty = str(generated.get("pack_stations") or self._sum_matching(it_devices, ["packing", "pack station"]))
        return_qty = str(generated.get("return_stations") or 0)
        other_qty = str(generated.get("other_stations") or 0)

        self._set_table_cell(doc, 1, 1, 1, admin_qty)
        self._set_table_cell(doc, 1, 2, 1, inbound_qty)
        self._set_table_cell(doc, 1, 3, 1, pack_qty)
        self._set_table_cell(doc, 1, 4, 1, return_qty)
        if len(doc.tables) > 1 and len(doc.tables[1].rows) > 5:
            self._set_table_cell(doc, 1, 5, 1, other_qty)

        # AI-generated scope narrative inserted after scope heading
        if generated.get("scope_narrative"):
            self._insert_after_heading(doc, "1. SCOPE OF WORK", generated["scope_narrative"])

        # Ensure deliverables section uses the AOT milestone-aligned standard
        self._apply_standard_deliverables(doc)

        # BoM appendix
        all_it = it_devices + classified_devices.get("Software", [])
        self._append_bom(doc, all_it, "IT Hardware – Bill of Materials")

    # ── OT Device ─────────────────────────────────────────────────────────────

    def _fill_ot_device(
        self,
        doc: Document,
        project_info: dict,
        classified_devices: dict,
        generated: dict,
    ) -> None:
        self._replace_all(doc, "Project Name: ", f"Project Name: {project_info.get('project_name', '')}")
        self._replace_all(doc, "FBM ID:", f"FBM ID: {project_info.get('fbm_id', '')}")
        self._replace_all(doc, "Site address: ", f"Site address: {project_info.get('site_address', '')}")
        self._replace_all(doc, "SITE ID: ", f"SITE ID: {project_info.get('site_id', '')}")
        self._replace_all(
            doc,
            "xx/xx/2026 – Access to Site",
            f"{project_info.get('access_date', 'TBD')} – Access to Site",
        )
        self._replace_all(
            doc,
            "xx/xx/2026 – Go live.",
            f"{project_info.get('go_live_date', 'TBD')} – Go live.",
        )

        self._set_table_cell(doc, 0, 1, 1, datetime.today().strftime("%d-%m-%Y"))

        if generated.get("scope_narrative"):
            self._insert_after_heading(doc, "SCOPE OF WORK", generated["scope_narrative"])

        # Ensure deliverables section uses the AOT milestone-aligned standard
        self._apply_standard_deliverables(doc)

        ot_devices = classified_devices.get("OT", [])
        self._append_bom(doc, ot_devices, "OT Hardware – Bill of Materials")

    # ── MDF ───────────────────────────────────────────────────────────────────

    def _fill_mdf(
        self,
        doc: Document,
        project_info: dict,
        classified_devices: dict,
        generated: dict,
    ) -> None:
        replacements: dict[str, str] = {
            "[Project Name]": project_info.get("project_name", ""),
            "[FBM ID]": project_info.get("fbm_id", ""),
            "[FbM ID]": project_info.get("fbm_id", ""),
            "[Site Address]": project_info.get("site_address", ""),
            "[Site ID]": project_info.get("site_id", ""),
            "[TBD – to be confirmed based on site headcount]": str(
                generated.get("users_supported", "TBD")
            ),
            "[TBD – confirm with total number of shifts]": "TBD",
            "[TBD – confirm total IDF count connected to this MDF]": str(
                generated.get("idf_count", "TBD")
            ),
            "[TBD – 1 / 2 / 3 rack configuration]": (
                str(generated.get("rack_count", "TBD")) + " rack(s)"
            ),
        }
        for old, new in replacements.items():
            self._replace_all(doc, old, new)

        # Issue date
        self._set_table_cell(doc, 0, 1, 1, datetime.today().strftime("%d/%m/%Y"))

        # Ensure deliverables section uses the AOT milestone-aligned standard
        self._apply_standard_deliverables(doc)

        net_devices = (
            classified_devices.get("Network", []) + classified_devices.get("MDF", [])
        )
        self._append_bom(doc, net_devices, "Network / MDF – Bill of Materials")

    # ── OT Automation Machine ─────────────────────────────────────────────────

    def _fill_ot_automation(
        self,
        doc: Document,
        project_info: dict,
        classified_devices: dict,
        generated: dict,
    ) -> None:
        replacements: dict[str, str] = {
            "[Project Name]": project_info.get("project_name", ""),
            "[FBM ID]": project_info.get("fbm_id", ""),
            "[FbM ID]": project_info.get("fbm_id", ""),
            "[Customer Name]": project_info.get("customer_name", ""),
            "[Site Address]": project_info.get("site_address", ""),
            "[Site ID]": project_info.get("site_id", ""),
        }
        # Replace [TBD] with go-live date where applicable
        go_live = project_info.get("go_live_date", "TBD")
        replacements["[TBD]"] = go_live

        for old, new in replacements.items():
            self._replace_all(doc, old, new)

        # Issue date
        self._set_table_cell(doc, 0, 1, 1, datetime.today().strftime("%d/%m/%Y"))

        auto_devices = classified_devices.get("Automation", [])
        if generated.get("scope_narrative"):
            self._insert_after_heading(
                doc, "OT Machine – Scope Overview", generated["scope_narrative"]
            )

        # Ensure deliverables section uses the AOT milestone-aligned standard
        self._apply_standard_deliverables(doc)

        self._append_bom(doc, auto_devices, "Automation Equipment – Bill of Materials")

    # ── Standard Deliverables ──────────────────────────────────────────────────

    def _apply_standard_deliverables(self, doc: Document) -> None:
        """
        Replace the Deliverables section in *doc* with the AOT milestone-aligned
        standard content defined in STANDARD_DELIVERABLES.

        The method locates the existing Deliverables heading (a paragraph whose
        heading style contains "Deliverable" as the last word, or whose heading
        text starts with a digit followed by "Deliverable"), removes everything
        between that heading and the next same-or-higher-level heading, then
        inserts the standardised sub-headings and bullet points.

        python-docx does not expose a public insert-after API, so paragraph
        elements are moved via their underlying lxml ``_element`` references.
        This is the standard approach recommended in the python-docx documentation
        for in-place structural modifications.

        This is called automatically when filling every template type so that
        the generated document always carries the canonical deliverables structure,
        regardless of what the source DOCX contained.
        """
        HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
        paragraphs = doc.paragraphs

        # Find the Deliverables section heading: must be a heading-style paragraph
        # whose normalised text ends with "deliverables" (case-insensitive).
        del_idx: int | None = None
        del_para = None
        for i, p in enumerate(paragraphs):
            if p.style.name in HEADING_STYLES:
                normalised = p.text.replace("\t", " ").strip().lower()
                if normalised.endswith("deliverables") or normalised == "deliverables":
                    del_idx = i
                    del_para = p
                    break
        if del_idx is None or del_para is None:
            return  # No deliverables heading found – nothing to standardise

        # Derive the section number (e.g. "4" from "4. Deliverables" or "5\tDeliverables")
        m = re.match(r"^(\d+)", del_para.text.replace("\t", " ").strip())
        sec_num = m.group(1) if m else _DEFAULT_DELIVERABLES_SECTION

        del_style = del_para.style.name  # e.g. "Heading 1" or "Heading 2"
        sub_style = "Heading 2" if del_style == "Heading 1" else "Heading 3"

        # Find the next major section heading at the same or higher level
        next_idx: int = len(paragraphs)
        major_styles = {"Heading 1", "Heading 2"} if del_style == "Heading 2" else {"Heading 1"}
        for i in range(del_idx + 1, len(paragraphs)):
            p = paragraphs[i]
            if p.style.name in major_styles and p.text.strip():
                next_idx = i
                break

        # Remove all paragraphs between the deliverables heading and next section.
        # python-docx does not provide a delete-paragraph API; direct lxml element
        # removal is the accepted pattern for structural document editing.
        for p in paragraphs[del_idx + 1 : next_idx]:
            p._element.getparent().remove(p._element)

        # `del_para` was not deleted; use it directly as the insertion anchor.
        anchor = del_para

        # Insert standardised milestone sub-sections immediately after the heading.
        for heading_tmpl, bullets in STANDARD_DELIVERABLES:
            h_para = doc.add_paragraph(heading_tmpl.format(sec=sec_num), style=sub_style)
            anchor._element.addnext(h_para._element)
            anchor = h_para
            for bullet in bullets:
                b_para = doc.add_paragraph(bullet, style="List Paragraph")
                anchor._element.addnext(b_para._element)
                anchor = b_para

    # ── Generic fallback ──────────────────────────────────────────────────────

    def _fill_generic(
        self,
        doc: Document,
        project_info: dict,
        classified_devices: dict,
        generated: dict,
    ) -> None:
        for key, value in project_info.items():
            if value:
                placeholder = f"[{key.replace('_', ' ').title()}]"
                self._replace_all(doc, placeholder, value)

    # ── DOCX utilities ─────────────────────────────────────────────────────────

    @staticmethod
    def _replace_in_paragraph(paragraph: Any, old: str, new: str) -> bool:
        """
        Replace *old* with *new* in a paragraph, merging runs to handle
        text that spans multiple runs.  Returns True if a replacement was made.
        """
        full_text = "".join(r.text for r in paragraph.runs)
        if old not in full_text:
            return False

        new_text = full_text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        return True

    def _replace_all(self, doc: Document, old: str, new: str) -> None:
        """Replace text in every paragraph and table cell of the document."""
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, old, new)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, old, new)

    @staticmethod
    def _set_table_cell(
        doc: Document,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        value: str,
    ) -> None:
        """Overwrite the text of a specific table cell."""
        try:
            cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            if para.runs:
                para.runs[0].text = value
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(value)
        except (IndexError, AttributeError):
            pass

    def _insert_after_heading(
        self,
        doc: Document,
        heading_text: str,
        new_text: str,
    ) -> None:
        """
        Insert a new paragraph with *new_text* immediately after the first
        paragraph whose text contains *heading_text*.
        """
        for idx, para in enumerate(doc.paragraphs):
            if heading_text.lower() in para.text.lower():
                # python-docx doesn't have an insert-after API; we use XML directly
                new_para = doc.add_paragraph(new_text)
                # Move the new paragraph's XML element to the right place
                para._element.addnext(new_para._element)
                return

    def _append_bom(
        self,
        doc: Document,
        devices: list[dict],
        title: str,
    ) -> None:
        """Append a Bill of Materials table at the end of the document."""
        if not devices:
            return

        doc.add_page_break()
        heading = doc.add_heading(title, level=2)
        doc.add_paragraph(
            f"Auto-generated on: {datetime.today().strftime('%d %B %Y')} "
            "(review and validate before final release)"
        )

        table = doc.add_table(rows=1, cols=4)
        # Use 'Table Grid' if available, otherwise fall back to the first table style
        try:
            table.style = "Table Grid"
        except KeyError:
            pass  # Keep the default style applied by python-docx

        hdr = table.rows[0].cells
        for i, label in enumerate(("#", "Description", "Quantity", "Category")):
            hdr[i].text = label
            # Bold header
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True

        for idx, device in enumerate(devices, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = device["name"]
            row[2].text = str(device["quantity"])
            row[3].text = device.get("category", "")

    # ── Quantity helper ───────────────────────────────────────────────────────

    @staticmethod
    def _sum_matching(devices: list[dict], keywords: list[str]) -> int:
        """Sum quantities of devices whose name contains any keyword."""
        return sum(
            d["quantity"]
            for d in devices
            if any(kw.lower() in d["name"].lower() for kw in keywords)
        )
